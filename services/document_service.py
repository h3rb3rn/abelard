"""Service: Upload, Text-Extraktion und Vektor-Indexierung von Debatten-Material.

Unterstuetzte Formate:
- Dokumente: .txt, .md, .pdf, .docx  → Text wird extrahiert, gechunkt und in
  ChromaDB (Collection ``project_documents``) indexiert.
- Bilder: .png, .jpg, .jpeg, .webp, .gif → Metadaten via Pillow; die optionale
  Benutzer-Beschreibung wird indexiert, damit das Bild im Debatten-Retrieval
  auffindbar ist.

Alle Index-Eintraege tragen ``project_id`` als Metadatum — das Retrieval im
Orchestrator ist strikt pro Projekt gescoped.
"""

from __future__ import annotations

import asyncio
import logging
import re
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import chromadb

from config import settings

logger = logging.getLogger(__name__)

DOCUMENT_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
ALLOWED_EXTENSIONS = DOCUMENT_EXTENSIONS | IMAGE_EXTENSIONS

CHUNK_SIZE_CHARS = 1200
CHUNK_OVERLAP_CHARS = 200
_COLLECTION_NAME = "project_documents"


class UploadValidationError(ValueError):
    """Ungueltige Datei (Typ oder Groesse)."""


@dataclass
class ExtractionResult:
    text: str
    kind: str  # "document" | "image"
    meta: dict[str, Any]


def safe_filename(filename: str) -> str:
    """Strip path components and dangerous characters from an upload filename."""
    name = Path(filename or "upload").name
    name = re.sub(r"[^\w.\-äöüÄÖÜß ]", "_", name).strip() or "upload"
    return name[:200]


def validate_upload(filename: str, size_bytes: int) -> str:
    """Return the file extension or raise UploadValidationError."""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise UploadValidationError(f"Dateityp '{ext}' nicht erlaubt. Erlaubt: {allowed}")
    if size_bytes > settings.upload_max_bytes:
        limit_mb = settings.upload_max_bytes // (1024 * 1024)
        raise UploadValidationError(f"Datei zu gross (max. {limit_mb} MB)")
    if size_bytes == 0:
        raise UploadValidationError("Leere Datei")
    return ext


def resolve_upload_root() -> Path:
    """Upload-Verzeichnis mit Fallback fuer lokale Entwicklung ohne /data."""
    root = Path(settings.upload_dir)
    try:
        root.mkdir(parents=True, exist_ok=True)
        return root
    except (PermissionError, OSError):
        fallback = Path("./data/uploads").resolve()
        fallback.mkdir(parents=True, exist_ok=True)
        logger.warning("Upload dir %s nicht beschreibbar — Fallback %s", settings.upload_dir, fallback)
        return fallback


def store_file(project_id: str, doc_id: str, filename: str, data: bytes) -> Path:
    """Persist raw upload bytes under {upload_root}/{project_id}/{doc_id}_{name}."""
    target_dir = resolve_upload_root() / project_id
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"{doc_id}_{safe_filename(filename)}"
    target.write_bytes(data)
    return target


# --------------------------------------------------------------------------- #
# Text-Extraktion                                                             #
# --------------------------------------------------------------------------- #

def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_image_meta(path: Path) -> dict[str, Any]:
    try:
        from PIL import Image

        with Image.open(path) as img:
            return {"width": img.width, "height": img.height, "format": img.format}
    except Exception as exc:
        logger.warning("Bild-Metadaten fehlgeschlagen fuer %s: %s", path.name, exc)
        return {}


def extract_content(path: Path, description: str = "") -> ExtractionResult:
    """Extract indexable text from an uploaded file (sync — call via to_thread)."""
    ext = path.suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        meta = _extract_image_meta(path)
        # Fuer Bilder ist die Beschreibung der indexierbare Text
        return ExtractionResult(text=description.strip(), kind="image", meta=meta)

    if ext == ".pdf":
        text = _extract_pdf(path)
    elif ext == ".docx":
        text = _extract_docx(path)
    else:  # .txt / .md
        text = path.read_bytes().decode("utf-8", errors="replace")

    if description.strip():
        text = f"{description.strip()}\n\n{text}"
    return ExtractionResult(text=text, kind="document", meta={})


def chunk_text(text: str, size: int = CHUNK_SIZE_CHARS, overlap: int = CHUNK_OVERLAP_CHARS) -> list[str]:
    """Split text into overlapping chunks along paragraph/sentence boundaries."""
    text = text.strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            # An Absatz- oder Satzgrenze zurueckschneiden, wenn moeglich
            window = text[start:end]
            cut = max(window.rfind("\n\n"), window.rfind(". "), window.rfind("! "), window.rfind("? "))
            if cut > size // 2:
                end = start + cut + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return chunks


# --------------------------------------------------------------------------- #
# Vektor-Index (ChromaDB)                                                     #
# --------------------------------------------------------------------------- #

class DocumentIndex:
    """Project-scoped vector index over uploaded material."""

    def __init__(self, persist_dir: str = "") -> None:
        self._client = chromadb.PersistentClient(path=persist_dir or settings.chroma_persist_dir)
        self._collection = self._client.get_or_create_collection(name=_COLLECTION_NAME)

    async def index_document(
        self,
        doc_id: str,
        project_id: str,
        filename: str,
        kind: str,
        text: str,
    ) -> int:
        """Chunk and index a document's text. Returns the number of chunks."""
        chunks = chunk_text(text)
        if not chunks:
            return 0
        ids = [f"{doc_id}:{i}" for i in range(len(chunks))]
        metas = [
            {"doc_id": doc_id, "project_id": project_id, "filename": filename, "kind": kind, "chunk": i}
            for i in range(len(chunks))
        ]
        await asyncio.to_thread(self._collection.add, ids=ids, documents=chunks, metadatas=metas)
        return len(chunks)

    async def remove_document(self, doc_id: str) -> None:
        try:
            await asyncio.to_thread(self._collection.delete, where={"doc_id": doc_id})
        except Exception as exc:
            logger.warning("DocumentIndex remove warning fuer %s: %s", doc_id, exc)

    async def search(self, project_id: str, query: str, top_k: int = 4) -> list[dict[str, Any]]:
        """Return the most relevant material chunks for a project."""
        try:
            results = await asyncio.to_thread(
                self._collection.query,
                query_texts=[query],
                n_results=top_k,
                where={"project_id": project_id},
            )
            records: list[dict[str, Any]] = []
            docs = (results.get("documents") or [[]])[0]
            metas = (results.get("metadatas") or [[]])[0]
            for doc, meta in zip(docs, metas):
                rec = dict(meta) if isinstance(meta, dict) else {}
                rec["document"] = doc
                records.append(rec)
            return records
        except Exception as exc:
            logger.warning("DocumentIndex search warning: %s", exc)
            return []


_index: DocumentIndex | None = None


def get_document_index() -> DocumentIndex:
    """Lazy singleton — eine ChromaDB-Instanz pro Prozess."""
    global _index
    if _index is None:
        _index = DocumentIndex()
    return _index
